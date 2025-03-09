# Copyright (c) 2025 Kyle Lopin (Naresuan University) <kylel@nu.ac.th>

"""
Module for generating student-specific Word documents with prediction datasets for Fish and Iris.

This module provides a function to create a Word document with formatted tables for Fish and Iris prediction
datasets. The datasets are generated dynamically, and columns for missing values (e.g., Weight, Target)
are left blank for students to fill.

Features:
- Generates prediction tables for Fish and Iris datasets.
- Allows formatting numeric values to one decimal place.
- Saves the output document with the student's name and a specified filename.

Usage:
    1. Prepare prediction datasets as pandas DataFrames.
    2. Call `create_document` with the student's name and the datasets.

"""

__author__ = "Kyle Vitautas Lopin"

# standard libraries
from io import BytesIO
from pathlib import Path
import pickle

# installed libraries
from docx import Document
from docx.shared import Inches
import pandas as pd

# local files
import get_data  # for testing
# Format Constants
MAX_COLUMN_WIDTH = Inches(1.5)


def insert_table_at_paragraph(doc, paragraph, df, extra_column_name=None):
    """
    Insert a table at the location of a paragraph and remove the paragraph.

    Args:
        doc (Document): The Word document object.
        paragraph (Paragraph): The paragraph to be replaced with the table.
        df (pd.DataFrame): The DataFrame to convert into a table.
        extra_column_name (str): Name of the extra column to add to the table.
    """
    # Ensure extra_column_name is a list
    if isinstance(extra_column_name, str):
        extra_column_name = [extra_column_name]  # Convert single name to list

    # Get the parent element of the paragraph
    parent_element = paragraph._element

    # Add a new table
    if extra_column_name:  # with an extra column for the students to fill in
        table = doc.add_table(rows=1, cols=len(df.columns) + len(extra_column_name))
    else:  # just put in the column
        table = doc.add_table(rows=1, cols=len(df.columns))
    table.autofit = False
    table.allow_autofit = False
    table.style = "Table Grid"

    # Add the header row
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
    if extra_column_name:
        for j, extra_col in enumerate(extra_column_name):
            hdr_cells[len(df.columns) + j].text = extra_col  # Add multiple extra column names

    # Add the dataset rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
        if extra_column_name:
            for j in range(len(extra_column_name)):  # Add empty cells for extra columns
                row_cells[len(df.columns) + j].text = ""

    for row in table.rows:
        for cell in row.cells:
            if cell.width > MAX_COLUMN_WIDTH:
                cell.width = MAX_COLUMN_WIDTH

    # Insert the table into the document at the correct location
    parent_element.addnext(table._element)

    # Remove the placeholder paragraph
    parent_element.getparent().remove(parent_element)


def create_document(student_name: str, output_file: str,
                    template_name: str,
                    tables: dict = {},
                    problems: list = [],
                    answers: list = [],
                    debug: bool = False) -> None:
    """
    Create a Word document for a student using a template, replacing placeholders
    with provided problem descriptions, answers, and tables.

    Args:
        student_name (str): The student's name.
        tables (dict): Dictionary of placeholders and DataFrames to replace them with.
        template_name (str): Name of the template in the pickle file.
        output_file (str): Name of the output Word document.
        problems (list): List of problem descriptions.
        answers (list): List of answers corresponding to each problem.
    """
    if debug:
        print(f"make document with parameters: template_name - {template_name}, "
              f"problems: {problems}")
    template_file = Path(__file__).parent / "templates" / "templates.pkl"
    # get document from templates
    with open(template_file, "rb") as file:
        templates = pickle.load(file)

    if template_name not in templates:
        print(f"template keys: {templates.keys()}")
        raise ValueError(f"Template '{template_name}' not found in templates.pkl.")

    doc = templates[template_name]

    # Load the binary datasets into a Document object, make it harder for students to read
    if isinstance(doc, bytes):
        doc = Document(BytesIO(doc))

    # replace student name
    for paragraph in doc.paragraphs:
        if "{Name}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{Name}", student_name)

    # Loop through problems and replace `{problem1}`, `{problem2}`, etc.
    for label, items in zip(["problem", "answer"], [problems, answers]):
        for i, item in enumerate(items):
            placeholder = f"{{{label}{i+1}}}"

            for paragraph in doc.paragraphs:
                paragraph.text = paragraph.text.replace(placeholder, item)

    # Replace table placeholders with actual tables
    for table_placeholder, df_n_extra_column_names in tables.items():
        for paragraph in doc.paragraphs:
            if table_placeholder in paragraph.text:
                # Replace the placeholder paragraph with the table
                if (isinstance(df_n_extra_column_names, list) and
                        len(df_n_extra_column_names) >= 2):
                    df, target_column_name = df_n_extra_column_names
                    insert_table_at_paragraph(doc, paragraph, df, target_column_name)
                else:
                    insert_table_at_paragraph(doc, paragraph, df_n_extra_column_names[0])
                break  # Stop searching after replacing the placeholder

    # Save the document
    doc.save(f"{student_name}_{output_file}")
    print(f"Document saved as: {student_name}_{output_file}")


def test_midterm_2025():
    # Example fish prediction datasets (excluding Weight)
    fish_prediction_data = pd.DataFrame({
        "Length": [30.1, 31.5, 29.8, 32.2, 30.7],
    })

    # Example iris prediction datasets (excluding Target)
    iris_prediction_data = pd.DataFrame({
        "sepal length (cm)": [5.1, 5.5, 6.0, 6.2, 5.8],
        "sepal width (cm)": [3.5, 3.8, 3.2, 3.0, 3.3],
        "petal length (cm)": [1.4, 1.5, 1.6, 1.7, 1.8],
        "petal width (cm)": [0.2, 0.3, 0.4, 0.5, 0.6]
    })

    # Student name
    _student_name = "John Doe"
    _tables = {"{Table 1}": [fish_prediction_data, "Predicted species"],
               "{Table 2}": [iris_prediction_data, "Predicted species"]}

    # Create the document
    create_document(_student_name, _tables)


if __name__ == '__main__':
    getdata = get_data.GetData(43)
    species2 = "Teak"
    df_trees, x_trees, prices = getdata.load_data(
        "Thai trees class and regr",
        num_points=300, test_size=10)
    df_trees.reset_index(drop=True, inplace=True)
    x_trees.reset_index(drop=True, inplace=True)

    # get data for later use in problem 3 and 4
    df_spectral, sensor, leaf = getdata.load_data("chloro")

    cost_table = {"{Table 1}": [x_trees, "Cost"]}

    questions = [(f"Fill in the table with price expected for each "
                  "tree below.", cost_table),
                 (f"For the tree species {species2}.\n"
                  "\nWhat is the average increase in weight for a "
                  "1 m increase in its length: ",),
                 (f"for {sensor} measurements on {leaf} leaves,",)]

    # make document
    # figure out what this take a table dict
    price_df = pd.DataFrame(list(prices.items()), columns=["Species", "Price"])
    student_name="Kyle"
    doc_table = {"{Table 1}": [pd.DataFrame(x_trees), "Cost"],
                 "{Table 2}": [price_df]}
    create_document(student_name=student_name,
                    output_file="Final_Spring_2025.docx",
                    template_name="Final_s25",
                    problems=[questions[0][0], questions[1][0],
                              questions[2][0]],
                    tables=doc_table)

    ham
    fish_prediction_data = pd.DataFrame({
        "Fish number": ["Fish 1", "Fish 2", "Fish 3"],
        "Length": [30.1, 31.5, 29.8],
    })
    df = pd.read_csv("hf://datasets/scikit-learn/Fish/Fish.csv")
    import random, get_fish_data, get_data
    data_gen = get_data.GetData(43)

    synthetic_x = data_gen.load_data("fish syn", num_points=6)
    print(synthetic_x)

    species2, _ = data_gen.load_data("fish coeff")

    weight_table = {"{Table 1}": [synthetic_x, ["Species", "Weight"]]}

    questions = [(f"Fill in the table with the species and weight of fish for each entry.",
                  weight_table), (f"For the fish species {species2}.\n"
                  "\nWhat is the average increase in weight for a "
                  "1 cm increase in its length: ")]


    doc_table = {"{Table 1}": [weight_table, ""]}
    # student_name = "Kyle"
    # make document
    create_document(student_name, "final_review_2025.docx",
                    "midterm_review_25",
                    problems=[questions[0][0], questions[1]],
                    tables=weight_table)
    ham
    # Create the document
    # create_document("Kyle", "test.docx",
    #                 "fish_regression_intro", _tables,
    #                 problems=[f"For the fish species Test.\n"
    #                            "\nWhat the average increase in weight for a "
    #                            "1 cm increase in its length: ",
    #                           f"What is the total cost for all of the Test_specis listed below, "
    #                           f"if the fish cost _prie baht per 100 grams"],
    #                 answers=["46", "938"])

    student_name = "Kyle"
    species2 = "Bream"
    weight_coefficient = 2
    species1 = "Pike"
    price = 100
    cost = 5




